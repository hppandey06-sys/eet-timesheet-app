"""
mer_report_generator.py
=======================
Auto-generates the 6-page A4 Monthly Engineering Review (MER) DOCX report
from live data in Supabase.

Same format as the canonical Meeting #04 report template, but populated
with real data for any selected meeting.

Sections produced (all A4 portrait):
  1. Cover + meeting metadata
  2. Key Highlights (timesheet KPIs)
  3. Project Status Table (from PREP submissions + timesheet)
  4. Manhours Booking (discipline-wise from timesheet)
  5. Engineering Strength (from members table)
  6. Critical Items (aggregated from PREP)
  7. Open Action Items (from ts_mer_actions)
  8. Interdisciplinary Coordination
  9. Software & Training
 10. Future / Closing notes

Usage:
  from mer_report_generator import generate_report
  buffer = generate_report(meeting_id)
  # buffer is BytesIO containing a .docx file
"""
from io import BytesIO
from datetime import date, datetime
import pandas as pd

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import db


# ════════════════════════════════════════════════════
# COLORS  (matching the canonical template)
# ════════════════════════════════════════════════════
NAVY = RGBColor(0x1F, 0x4E, 0x78)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
LIGHT_BG = "F2F7FB"
HEADER_BG = "1F4E78"
ROW_ALT = "F8F9FB"
RED = RGBColor(0xC0, 0x00, 0x00)
AMBER = RGBColor(0xBF, 0x8F, 0x00)
GREEN = RGBColor(0x0F, 0x6E, 0x56)
GREY = RGBColor(0x59, 0x59, 0x59)


# ════════════════════════════════════════════════════
# HELPERS
# ════════════════════════════════════════════════════

def _shade_cell(cell, color_hex):
    """Add shading to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _add_run(para, text, font="Calibri", size=11, bold=False, color=None, italic=False):
    """Add a styled text run to a paragraph."""
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def _add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    _add_run(p, text, size=16, bold=True, color=NAVY)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_para(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    _add_run(p, text, size=size, bold=bold, color=color)
    return p


def _add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    _add_run(p, text, size=size)
    return p


def _make_table(doc, n_rows, n_cols, col_widths_inches=None):
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths_inches:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths_inches):
                    cell.width = Inches(col_widths_inches[i])
    return table


def _header_cell(cell, text):
    """Style as table header."""
    _shade_cell(cell, HEADER_BG)
    p = cell.paragraphs[0]
    _add_run(p, text, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _data_cell(cell, text, bold=False, color=None, align=None, size=10, fill=None):
    if fill:
        _shade_cell(cell, fill)
    p = cell.paragraphs[0]
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, text or '', size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


# ════════════════════════════════════════════════════
# DATA LOADERS
# ════════════════════════════════════════════════════

def _load_meeting(meeting_id):
    sb = db.get_client()
    r = sb.table('ts_mer_meetings').select('*').eq('id', meeting_id).execute()
    return r.data[0] if r.data else None


def _load_inputs(meeting_id):
    sb = db.get_client()
    r = sb.table('ts_mer_inputs').select('*').eq('meeting_id', meeting_id).execute()
    return r.data or []


def _load_actions(include_closed=False):
    sb = db.get_client()
    statuses = ['OPEN', 'IN_PROGRESS'] if not include_closed else ['OPEN', 'IN_PROGRESS', 'CLOSED']
    r = sb.table('ts_mer_actions').select('*').in_('status', statuses).order('due_date').execute()
    return r.data or []


def _load_decisions(meeting_id):
    sb = db.get_client()
    r = sb.table('ts_mer_decisions').select('*').eq('meeting_id', meeting_id).execute()
    return r.data or []


def _load_attendance(meeting_id):
    sb = db.get_client()
    r = sb.table('ts_mer_attendance').select('*').eq('meeting_id', meeting_id).execute()
    return r.data or []


def _load_entries_for_month(year, month):
    sb = db.get_client()
    start = date(year, month, 1).isoformat()
    if month == 12:
        end = date(year + 1, 1, 1).isoformat()
    else:
        end = date(year, month + 1, 1).isoformat()
    all_rows = []
    offset = 0
    while True:
        r = (sb.table('ts_entries')
               .select('*')
               .gte('entry_date', start)
               .lt('entry_date', end)
               .order('entry_date')
               .order('id')
               .range(offset, offset + 999)
               .execute())
        batch = r.data or []
        all_rows.extend(batch)
        if len(batch) < 1000:
            break
        offset += 1000
    return all_rows


# ════════════════════════════════════════════════════
# REPORT BUILDER SECTIONS
# ════════════════════════════════════════════════════

def _build_cover(doc, meeting):
    """Add title page + meeting metadata table."""
    review_month = meeting['review_month']
    year, month = map(int, review_month.split('-'))
    month_name = date(year, month, 1).strftime('%B %Y')
    last_day = date(year, month + 1, 1) if month < 12 else date(year + 1, 1, 1)
    last_day = last_day.replace(day=1)
    last_day_actual = (last_day - pd.Timedelta(days=1)).date() if month < 12 else date(year, 12, 31)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    _add_run(p, "GCC ENGINEERING", size=22, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(p, "Monthly Review Meeting — Minutes & Action Tracker", size=14, color=GREY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    _add_run(p, f"Meeting #{meeting['meeting_no']:02d}  ·  Reviewing {month_name}",
             size=18, bold=True, color=ACCENT)

    # Metadata table
    table = doc.add_table(rows=5, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cells = [
        ("Meeting:", "GCC Engineering Monthly Review",
         "Meeting No:", f"{meeting['meeting_no']:02d}"),
        ("Reviewing:", f"{month_name} (data locked {last_day_actual.strftime('%d-%b-%Y')})",
         "Held on:", meeting.get('scheduled_date', 'TBD') or 'TBD'),
        ("Location:", "GCC Arioli, Mumbai",
         "Status:", meeting.get('status', 'DRAFT')),
        ("Chair / Presenter:", meeting.get('chair', 'Hariprakash Pandey'),
         "Reviewer:", meeting.get('reviewer', 'Sangeeta Salvi')),
        ("Attendees:", "", "", ""),
    ]

    for i, row_data in enumerate(cells):
        row = table.rows[i]
        if i < 4:
            for j, txt in enumerate(row_data):
                bold = (j % 2 == 0)
                fill = LIGHT_BG if bold else None
                color = NAVY if (i == 0 and j == 3) else None
                _data_cell(row.cells[j], txt, bold=bold, fill=fill, color=color, size=10)
        else:
            # Attendees row spans cols 1-3
            _data_cell(row.cells[0], "Attendees:", bold=True, fill=LIGHT_BG, size=10)
            # Merge cells 1,2,3
            merged = row.cells[1].merge(row.cells[3])
            attendance = _load_attendance(meeting['id'])
            present_count = sum(1 for a in attendance if a.get('status') == 'PRESENT')
            if attendance:
                _data_cell(merged, f"{present_count} of {len(attendance)} attended", size=10)
            else:
                _data_cell(merged, "Engineering Team (attendance recorded in app)", size=10)

    doc.add_paragraph()


def _build_key_highlights(doc, meeting, entries_df, members):
    _add_heading1(doc, "1. Key Highlights")
    _add_para(doc, f"Auto-generated from timesheet · data locked end of {meeting['review_month']}",
              size=9, color=GREY)

    if len(entries_df) == 0:
        _add_para(doc, "No timesheet entries for this period.")
        return

    excluded_uids = {m['id'] for m in members if m.get('excluded_from_productivity')}
    df_eng = entries_df[~entries_df['uid'].isin(excluded_uids)]

    total_hrs = df_eng['hrs'].sum()
    n_reporting = df_eng['uid'].nunique()
    eng_total = sum(1 for m in members
                    if m.get('dept') == 'Engineering' and m['id'] not in excluded_uids)
    n_projects = df_eng['proj'].nunique()

    top_proj_g = df_eng.groupby('proj')['hrs'].sum().sort_values(ascending=False)
    top_proj = top_proj_g.index[0] if len(top_proj_g) else '-'
    top_proj_hrs = top_proj_g.iloc[0] if len(top_proj_g) else 0
    top_proj_pct = top_proj_hrs / total_hrs * 100 if total_hrs else 0

    # 3-column KPI table
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    kpis = [
        ("TOTAL HOURS BOOKED", f"{total_hrs:,.0f}",
         f"by {n_reporting} of {eng_total} Engineering team ({n_reporting/max(eng_total,1)*100:.0f}%)"),
        ("PROJECTS ACTIVE", str(n_projects), "plus enabling work / overheads"),
        ("TOP PROJECT (by hrs)", str(top_proj), f"{top_proj_hrs:.0f} hrs · {top_proj_pct:.0f}% of total"),
    ]

    for i, (label, value, foot) in enumerate(kpis):
        cell = table.rows[0].cells[i]
        _shade_cell(cell, LIGHT_BG)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, label, size=8, color=GREY)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, value, size=20, bold=True, color=NAVY)

        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, foot, size=9, color=GREY)

    doc.add_paragraph()


def _build_project_status(doc, meeting, entries_df, inputs):
    _add_heading1(doc, "2. Project Status — Key Updates")
    _add_para(doc, "Aggregated from discipline lead inputs · Hours from timesheet",
              size=9, color=GREY)

    # Hours by project from timesheet
    if len(entries_df):
        proj_hours = entries_df.groupby('proj')['hrs'].sum().to_dict()
    else:
        proj_hours = {}

    # Project status from PREP submissions (JSONB)
    proj_statuses = {}
    for inp in inputs:
        ps = inp.get('project_status') or []
        if isinstance(ps, list):
            for p in ps:
                pname = (p.get('project') or '').strip()
                if pname:
                    if pname not in proj_statuses:
                        proj_statuses[pname] = []
                    proj_statuses[pname].append({
                        'phase': p.get('phase', ''),
                        'pct': p.get('pct', ''),
                        'status': p.get('status', ''),
                        'discipline': inp.get('discipline', ''),
                    })

    # Union of all projects
    all_projects = set(proj_hours.keys()) | set(proj_statuses.keys())
    # Sort by hours descending
    sorted_projects = sorted(all_projects,
                             key=lambda p: -proj_hours.get(p, 0))

    if not sorted_projects:
        _add_para(doc, "No projects with activity this month.")
        return

    table = doc.add_table(rows=len(sorted_projects) + 1, cols=5)
    headers = ['Sl.', 'Project', 'Phase', 'Apr Hrs', 'Status / Activities']
    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h)

    for idx, pname in enumerate(sorted_projects, 1):
        row = table.rows[idx]
        hours = proj_hours.get(pname, 0)
        statuses = proj_statuses.get(pname, [])
        phase = statuses[0]['phase'] if statuses else ''
        status_combined = ' / '.join(s['status'] for s in statuses if s.get('status')) if statuses else '— No PREP input yet —'

        fill = ROW_ALT if idx % 2 == 0 else None
        _data_cell(row.cells[0], str(idx), align='center', fill=fill)
        _data_cell(row.cells[1], pname, bold=True, fill=fill)
        _data_cell(row.cells[2], phase or '—', fill=fill)
        _data_cell(row.cells[3], f"{hours:.0f}" if hours else '—', align='right', bold=True, fill=fill)
        _data_cell(row.cells[4], status_combined[:300], fill=fill)


def _build_manhours(doc, entries_df, members):
    _add_heading1(doc, "3. Manhours Booking — Discipline-wise")
    _add_para(doc,
              "Manhours recorded diligently — essential for project claims. Demonstrates GCC Team value.",
              size=10)

    if len(entries_df) == 0:
        _add_para(doc, "No data.")
        return

    # Map uid → discipline
    mem_lookup = {m['id']: m for m in members}
    entries_df = entries_df.copy()
    entries_df['discipline'] = entries_df['uid'].map(
        lambda u: mem_lookup.get(u, {}).get('discipline', '-')
    )

    # Group
    excluded_uids = {m['id'] for m in members if m.get('excluded_from_productivity')}
    df_eng = entries_df[~entries_df['uid'].isin(excluded_uids)]
    df_mgmt = entries_df[entries_df['uid'].isin(excluded_uids)]

    disc_summary = (df_eng.groupby('discipline')
                          .agg(Hours=('hrs', 'sum'), Members=('uid', 'nunique'))
                          .reset_index()
                          .sort_values('Hours', ascending=False))

    total_eng_hrs = df_eng['hrs'].sum()
    mgmt_hrs = df_mgmt['hrs'].sum()
    grand_total = total_eng_hrs + mgmt_hrs

    n_rows = len(disc_summary) + 3  # + management row + total row + header
    table = doc.add_table(rows=n_rows, cols=4)
    headers = ['Discipline', 'Members', 'Hours', '% of Total']
    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h)

    for idx, (_, r) in enumerate(disc_summary.iterrows(), 1):
        row = table.rows[idx]
        pct = r['Hours'] / grand_total * 100 if grand_total else 0
        bars = '━' * round(pct / 2)
        fill = ROW_ALT if idx % 2 == 0 else None
        _data_cell(row.cells[0], r['discipline'], bold=True, fill=fill)
        _data_cell(row.cells[1], str(int(r['Members'])), align='center', fill=fill)
        _data_cell(row.cells[2], f"{r['Hours']:.0f}", align='right', bold=True, fill=fill)
        _data_cell(row.cells[3], f"{pct:.0f}%  {bars}", fill=fill)

    # Management row (excluded from productivity)
    mgmt_row = table.rows[len(disc_summary) + 1]
    n_mgmt_members = df_mgmt['uid'].nunique()
    _data_cell(mgmt_row.cells[0], "Engineering Management (HP+Sangeeta — excluded)",
               bold=True, fill=LIGHT_BG)
    _data_cell(mgmt_row.cells[1], str(n_mgmt_members), align='center', fill=LIGHT_BG)
    _data_cell(mgmt_row.cells[2], f"{mgmt_hrs:.0f}", align='right', bold=True, fill=LIGHT_BG)
    _data_cell(mgmt_row.cells[3], "Excluded from productivity scoring", fill=LIGHT_BG)

    # Total row
    total_row = table.rows[-1]
    for i in range(4):
        _shade_cell(total_row.cells[i], "1F4E78")
    n_total = df_eng['uid'].nunique() + n_mgmt_members
    _data_cell(total_row.cells[0], "TOTAL", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    _data_cell(total_row.cells[1], str(n_total), align='center', bold=True,
               color=RGBColor(0xFF, 0xFF, 0xFF))
    _data_cell(total_row.cells[2], f"{grand_total:.0f}", align='right', bold=True,
               color=RGBColor(0xFF, 0xFF, 0xFF))
    _data_cell(total_row.cells[3], "100%", bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def _build_engineering_strength(doc, members):
    _add_heading1(doc, "4. Engineering Strength — Discipline-wise")

    eng_members = [m for m in members if m.get('dept') == 'Engineering']
    disc_counts = {}
    for m in eng_members:
        d = m.get('discipline', '-') or '-'
        disc_counts[d] = disc_counts.get(d, 0) + 1

    total_strength = sum(disc_counts.values())
    _add_para(doc, f"Total Engineering strength: {total_strength} members.", size=10)

    rows_data = []
    for d, count in sorted(disc_counts.items(), key=lambda x: -x[1]):
        lead = next(
            (m['name'] for m in members
             if m.get('is_discipline_lead') and m.get('leads_discipline') == d),
            '—'
        )
        rows_data.append((d, count, lead))

    table = doc.add_table(rows=len(rows_data) + 2, cols=3)
    headers = ['Discipline', 'Strength', 'Discipline Lead']
    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h)

    for idx, (d, count, lead) in enumerate(rows_data, 1):
        row = table.rows[idx]
        fill = ROW_ALT if idx % 2 == 0 else None
        _data_cell(row.cells[0], d, fill=fill)
        _data_cell(row.cells[1], str(count), align='center', bold=True, fill=fill)
        _data_cell(row.cells[2], lead, fill=fill)

    total_row = table.rows[-1]
    for i in range(3):
        _shade_cell(total_row.cells[i], "1F4E78")
    _data_cell(total_row.cells[0], "TOTAL ENGINEERING", bold=True,
               color=RGBColor(0xFF, 0xFF, 0xFF))
    _data_cell(total_row.cells[1], str(total_strength), align='center', bold=True,
               color=RGBColor(0xFF, 0xFF, 0xFF))
    _data_cell(total_row.cells[2], "All disciplines have nominated leads", bold=True,
               color=RGBColor(0xFF, 0xFF, 0xFF))


def _build_critical_items(doc, inputs):
    _add_heading1(doc, "5. Critical Items & Decisions Needed")
    _add_para(doc,
              "Items requiring Engineering Head attention or escalation, "
              "aggregated from each discipline's PREP submissions.",
              size=10)

    has_items = False
    for inp in inputs:
        if inp.get('critical_items'):
            has_items = True
            p = doc.add_paragraph()
            _add_run(p, f"From {inp.get('discipline', '?')}:", size=11, bold=True, color=NAVY)
            doc.add_paragraph(inp['critical_items'])

    if not has_items:
        _add_para(doc,
                  "(No critical items submitted yet. Will be populated as discipline leads complete PREP.)",
                  size=10, color=GREY)


def _build_actions(doc, actions):
    _add_heading1(doc, "6. Open Action Items")
    _add_para(doc, "All actions currently open or in progress", size=9, color=GREY)

    if not actions:
        _add_para(doc, "(No open action items.)", size=10, color=GREY)
        return

    table = doc.add_table(rows=len(actions) + 1, cols=5)
    headers = ['Sl.', 'Action Item', 'Owner', 'Due', 'Status']
    for i, h in enumerate(headers):
        _header_cell(table.rows[0].cells[i], h)

    for idx, a in enumerate(actions, 1):
        row = table.rows[idx]
        fill = ROW_ALT if idx % 2 == 0 else None
        status_label = a.get('status', 'OPEN')
        status_color = GREEN if status_label == 'CLOSED' else (AMBER if status_label == 'IN_PROGRESS' else None)
        status_icon = '✅' if status_label == 'CLOSED' else ('🟡' if status_label == 'IN_PROGRESS' else '🔴')
        _data_cell(row.cells[0], str(idx), align='center', fill=fill)
        _data_cell(row.cells[1], a.get('action_text', ''), fill=fill)
        _data_cell(row.cells[2], a.get('owner_name', '—'), fill=fill)
        _data_cell(row.cells[3], a.get('due_date', '—') or '—', align='center', fill=fill)
        _data_cell(row.cells[4], f"{status_icon} {status_label}",
                   color=status_color, fill=fill)


def _build_coordination(doc, inputs):
    _add_heading1(doc, "7. Major Issues & Interdisciplinary Coordination")
    
    # Aggregate concerns from inputs
    concerns_by_disc = {}
    manpower_by_disc = {}
    for inp in inputs:
        if inp.get('concerns'):
            concerns_by_disc[inp.get('discipline', '?')] = inp['concerns']
        if inp.get('manpower_issues'):
            manpower_by_disc[inp.get('discipline', '?')] = inp['manpower_issues']

    if concerns_by_disc:
        _add_para(doc, "Concerns raised by disciplines:", size=11, bold=True)
        for disc, c in concerns_by_disc.items():
            p = doc.add_paragraph()
            _add_run(p, f"{disc}: ", size=10, bold=True)
            _add_run(p, c, size=10)
    
    if manpower_by_disc:
        _add_para(doc, "Manpower / resource issues:", size=11, bold=True)
        for disc, c in manpower_by_disc.items():
            p = doc.add_paragraph()
            _add_run(p, f"{disc}: ", size=10, bold=True)
            _add_run(p, c, size=10)

    if not concerns_by_disc and not manpower_by_disc:
        _add_para(doc,
                  "(Will be populated from PREP submissions and meeting discussions.)",
                  size=10, color=GREY)


def _build_software_training(doc, inputs):
    _add_heading1(doc, "8. Software & Training Needs")
    
    needs_by_disc = {}
    for inp in inputs:
        if inp.get('software_needs'):
            needs_by_disc[inp.get('discipline', '?')] = inp['software_needs']

    if needs_by_disc:
        for disc, c in needs_by_disc.items():
            p = doc.add_paragraph()
            _add_run(p, f"{disc}: ", size=10, bold=True)
            _add_run(p, c, size=10)
    else:
        _add_para(doc,
                  "(Will be populated from PREP submissions.)",
                  size=10, color=GREY)


def _build_decisions(doc, decisions):
    if not decisions:
        return
    _add_heading1(doc, "9. Decisions Captured")
    
    for d in decisions:
        p = doc.add_paragraph()
        _add_run(p, f"{d.get('topic', '?')}: ", size=11, bold=True, color=NAVY)
        _add_run(p, d.get('decision_text', ''), size=11)


def _build_closing(doc):
    _add_heading1(doc, "10. Notes for Discipline Leads (next meeting prep)")
    _add_para(doc,
              "From the next meeting onwards, this report is auto-generated by the MER module:",
              size=10)
    _add_bullet(doc, "1. PREP window opens on the 4th of the month (after timesheet locks 3rd EOD)", size=10)
    _add_bullet(doc, "2. Each discipline lead logs in and fills their section", size=10)
    _add_bullet(doc, "3. HP consolidates submissions and captures decisions during the meeting", size=10)
    _add_bullet(doc, "4. Sangeeta reviews the auto-generated report", size=10)
    _add_bullet(doc, "5. HP distributes the published version to the team", size=10)
    _add_para(doc, "Reference: gcc-eet-timesheet.streamlit.app → MER tab",
              size=9, color=GREY)


# ════════════════════════════════════════════════════
# MAIN ENTRY POINT
# ════════════════════════════════════════════════════

def generate_report(meeting_id):
    """
    Generate the 6-page A4 MER report DOCX for the given meeting.
    
    Returns: BytesIO buffer containing the .docx file
    """
    meeting = _load_meeting(meeting_id)
    if not meeting:
        raise ValueError(f"Meeting {meeting_id} not found")

    year, month = map(int, meeting['review_month'].split('-'))
    entries = _load_entries_for_month(year, month)
    members = db.get_members()
    inputs = _load_inputs(meeting_id)
    actions = _load_actions()
    decisions = _load_decisions(meeting_id)

    df = pd.DataFrame(entries) if entries else pd.DataFrame(columns=['uid', 'proj', 'hrs', 'description'])
    if len(df):
        df['hrs'] = df['hrs'].astype(float)

    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Header & footer
    header = doc.sections[0].header
    h_para = header.paragraphs[0]
    h_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_run(h_para,
             f"GCC Engineering Monthly Review · Meeting #{meeting['meeting_no']:02d} · "
             f"Reviewing {meeting['review_month']}",
             size=9, color=GREY)

    footer = doc.sections[0].footer
    f_para = footer.paragraphs[0]
    f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_run(f_para,
             f"Prepared by {meeting.get('chair', 'HP')} · "
             f"Reviewed by {meeting.get('reviewer', 'Sangeeta')} · "
             f"Generated {datetime.utcnow().strftime('%d-%b-%Y %H:%M')} UTC",
             size=9, color=GREY)

    # Build all sections
    _build_cover(doc, meeting)
    _build_key_highlights(doc, meeting, df, members)
    _build_project_status(doc, meeting, df, inputs)
    _build_manhours(doc, df, members)
    _build_engineering_strength(doc, members)
    _build_critical_items(doc, inputs)
    _build_actions(doc, actions)
    _build_coordination(doc, inputs)
    _build_software_training(doc, inputs)
    _build_decisions(doc, decisions)
    _build_closing(doc)

    # Save to buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_report_filename(meeting):
    """Standard filename for the generated report."""
    return (f"GCC_MER_Meeting_{meeting['meeting_no']:02d}_"
            f"{meeting['review_month'].replace('-', '_')}.docx")

